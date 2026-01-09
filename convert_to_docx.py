"""
Convert CPPS Final Report Markdown to Word Document (.docx)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_docx():
    input_file = 'CPPS_Final_Report.md'
    output_file = 'CPPS_Final_Report.docx'

    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found")
        return

    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set up page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Define styles
    styles = doc.styles

    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.name = 'Arial'
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)

    # Subtitle style
    subtitle_style = styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.size = Pt(14)
    subtitle_style.font.italic = True
    subtitle_style.font.name = 'Arial'
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(6)

    # Meta style (for title page info)
    meta_style = styles.add_style('CustomMeta', WD_STYLE_TYPE.PARAGRAPH)
    meta_style.font.size = Pt(11)
    meta_style.font.name = 'Arial'
    meta_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_style.paragraph_format.space_after = Pt(4)

    # Heading 1 style
    h1_style = styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.name = 'Arial'
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(10)

    # Heading 2 style
    h2_style = styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.size = Pt(13)
    h2_style.font.bold = True
    h2_style.font.name = 'Arial'
    h2_style.paragraph_format.space_before = Pt(14)
    h2_style.paragraph_format.space_after = Pt(8)

    # Heading 3 style
    h3_style = styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.size = Pt(11)
    h3_style.font.bold = True
    h3_style.font.name = 'Arial'
    h3_style.paragraph_format.space_before = Pt(10)
    h3_style.paragraph_format.space_after = Pt(6)

    # Body style
    body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.size = Pt(11)
    body_style.font.name = 'Arial'
    body_style.paragraph_format.space_after = Pt(8)
    body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Caption style
    caption_style = styles.add_style('CustomCaption', WD_STYLE_TYPE.PARAGRAPH)
    caption_style.font.size = Pt(10)
    caption_style.font.italic = True
    caption_style.font.name = 'Arial'
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.space_before = Pt(6)
    caption_style.paragraph_format.space_after = Pt(12)

    # Parse content
    lines = content.split('\n')
    in_title_section = True
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Main title
        if stripped.startswith('# ') and not stripped.startswith('## '):
            doc.add_paragraph()  # Add space at top
            doc.add_paragraph()
            title_text = stripped[2:].strip()
            p = doc.add_paragraph(title_text, style='CustomTitle')

        # Subtitle (## A Cyber-Physical...)
        elif in_title_section and stripped.startswith('## '):
            subtitle_text = stripped[3:].strip()
            doc.add_paragraph(subtitle_text, style='CustomSubtitle')
            doc.add_paragraph()  # Add space

        # Meta info on title page
        elif in_title_section and stripped.startswith('**') and ':**' in stripped:
            meta_text = stripped.replace('**', '')
            doc.add_paragraph(meta_text, style='CustomMeta')

        # --- creates page break
        elif stripped == '---':
            if in_title_section:
                in_title_section = False
            doc.add_page_break()

        # Section headers (## )
        elif stripped.startswith('## ') and not in_title_section:
            header_text = stripped[3:].strip()
            doc.add_paragraph(header_text, style='CustomH1')

        # Subsection headers (### )
        elif stripped.startswith('### '):
            header_text = stripped[4:].strip()
            doc.add_paragraph(header_text, style='CustomH2')

        # Sub-subsection headers (#### )
        elif stripped.startswith('#### '):
            header_text = stripped[5:].strip()
            doc.add_paragraph(header_text, style='CustomH3')

        # Screenshot placeholders
        elif stripped.startswith('[INSERT SCREENSHOT:'):
            caption = stripped.replace('[INSERT SCREENSHOT:', '').replace(']', '').strip()
            # Add placeholder box
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('[Insert Screenshot Here]')
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = None  # Gray color
            # Add caption
            doc.add_paragraph(f"Figure: {caption}", style='CustomCaption')

        # Bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            bullet_text = stripped[2:].strip()
            bullet_text = format_text_for_docx(bullet_text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, bullet_text)

        # Numbered lists
        elif re.match(r'^\d+\.', stripped):
            list_text = format_text_for_docx(stripped)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, list_text)

        # Regular paragraphs
        elif stripped and not stripped.startswith('#') and not stripped.startswith('```') and not stripped.startswith('|'):
            para_text = format_text_for_docx(stripped)
            p = doc.add_paragraph(style='CustomBody')
            add_formatted_text(p, para_text)

        i += 1

    # Save document
    doc.save(output_file)
    print(f"Word document created successfully: {output_file}")

def format_text_for_docx(text):
    """Keep markdown formatting markers for later processing"""
    # Remove markdown links, keep text
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text

def add_formatted_text(paragraph, text):
    """Add text with bold/italic formatting to paragraph"""
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            inner = part[2:-2]
            # Check for italic inside bold
            italic_parts = re.split(r'(\*.*?\*)', inner)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*'):
                    run = paragraph.add_run(ip[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    run = paragraph.add_run(ip)
                    run.bold = True
        else:
            # Check for italic
            italic_parts = re.split(r'(\*.*?\*)', part)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*') and not ip.startswith('**'):
                    run = paragraph.add_run(ip[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(ip)

if __name__ == "__main__":
    create_docx()
